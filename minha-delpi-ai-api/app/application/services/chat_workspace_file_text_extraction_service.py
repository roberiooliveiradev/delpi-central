"""Extração de texto de arquivos de workspace — pipeline canônico (Playbook 17).

Consumido por anexos de sessão, fontes de projeto/agente, contexto efêmero e admin.
PDF delega a ChatPdfDocumentExtractionService; imagens usam OCR opcional.
"""

from __future__ import annotations

import csv
import io
import json
import shutil
import subprocess
from pathlib import Path
from typing import Any

from app.application.services.chat_attachment_image_ocr_service import (
    ChatAttachmentImageOcrService,
)
from app.domain.services.chat_attachment_content_service import (
    ChatAttachmentContentService,
)


class ChatWorkspaceFileTextExtractionService:
    SUPPORTED_TEXT_EXTENSIONS = frozenset({".txt", ".md", ".markdown", ".csv", ".json"})
    SUPPORTED_OFFICE_EXTENSIONS = frozenset({".docx", ".xlsx", ".doc", ".xls"})
    SUPPORTED_DOCUMENT_EXTENSIONS = frozenset({".pdf"})
    SUPPORTED_IMAGE_EXTENSIONS = frozenset({".png", ".jpg", ".jpeg", ".webp"})

    @classmethod
    def supported_extensions(cls) -> set[str]:
        return (
            set(cls.SUPPORTED_TEXT_EXTENSIONS)
            | set(cls.SUPPORTED_OFFICE_EXTENSIONS)
            | set(cls.SUPPORTED_DOCUMENT_EXTENSIONS)
            | set(cls.SUPPORTED_IMAGE_EXTENSIONS)
        )

    @classmethod
    def extract(
        cls,
        *,
        storage_path: str,
        filename: str,
        content_type: str | None = None,
        pdf_page_limit: int | None = None,
    ) -> dict[str, Any]:
        path = Path(storage_path)
        extension = path.suffix.lower() or Path(filename).suffix.lower()

        if extension in cls.SUPPORTED_TEXT_EXTENSIONS:
            if extension == ".json":
                return cls._extract_json(path)

            if extension == ".csv":
                return cls._extract_csv(path)

            return cls._extract_plain_text(path, extension)

        if extension == ".doc":
            return cls._extract_doc(path)

        if extension == ".xls":
            return cls._extract_xls(path)

        if extension == ".docx":
            return cls._extract_docx(path)

        if extension == ".xlsx":
            return cls._extract_xlsx(path)

        if extension == ".pdf":
            return cls._extract_pdf(path, page_limit=pdf_page_limit)

        if extension in cls.SUPPORTED_IMAGE_EXTENSIONS:
            return cls._extract_image(path, extension)

        return {
            "supported": False,
            "content": "",
            "metadata": {
                "reason": "unsupported_extension",
                "extension": extension,
                "contentType": content_type,
            },
        }

    @classmethod
    def _extract_plain_text(cls, path: Path, extension: str) -> dict[str, Any]:
        content = path.read_text(encoding="utf-8", errors="ignore")

        return {
            "supported": True,
            "content": content,
            "metadata": {
                "extractor": "plain_text",
                "extension": extension,
            },
        }

    @classmethod
    def _extract_json(cls, path: Path) -> dict[str, Any]:
        raw = path.read_text(encoding="utf-8", errors="ignore")

        try:
            parsed = json.loads(raw)
            content = json.dumps(parsed, ensure_ascii=False, indent=2)
        except Exception:
            content = raw

        return {
            "supported": True,
            "content": content,
            "metadata": {
                "extractor": "json",
                "extension": ".json",
            },
        }

    @classmethod
    def _extract_csv(cls, path: Path) -> dict[str, Any]:
        raw = path.read_text(encoding="utf-8", errors="ignore")
        sample = raw[:4096]
        delimiter = ";" if sample.count(";") > sample.count(",") else ","
        reader = csv.reader(io.StringIO(raw), delimiter=delimiter)
        row_limit = ChatAttachmentContentService.file_extraction_csv_max_rows()

        rows: list[str] = []

        for index, row in enumerate(reader):
            if index >= row_limit:
                break

            rows.append(" | ".join(str(cell) for cell in row))

        return {
            "supported": True,
            "content": "\n".join(rows),
            "metadata": {
                "extractor": "csv",
                "extension": ".csv",
                "rowLimit": row_limit,
            },
        }

    @classmethod
    def _extract_docx(cls, path: Path) -> dict[str, Any]:
        try:
            from docx import Document
        except ImportError as exc:
            return cls._unsupported_optional(".docx", exc)

        try:
            document = Document(str(path))
        except Exception as exc:
            return cls._unsupported_optional(".docx", exc)

        lines: list[str] = []

        for paragraph in document.paragraphs:
            text = str(paragraph.text or "").strip()

            if text:
                lines.append(text)

        for table_index, table in enumerate(document.tables):
            lines.append(f"# Tabela {table_index + 1}")

            for row in table.rows:
                cells = [str(cell.text or "").strip() for cell in row.cells]

                if any(cells):
                    lines.append(" | ".join(cells))

        content = "\n".join(lines).strip()

        if not content:
            return cls._unsupported_optional(
                ".docx",
                ValueError("empty_docx_content"),
            )

        return {
            "supported": True,
            "content": content,
            "metadata": {
                "extractor": "python_docx",
                "extension": ".docx",
                "paragraphCount": len(document.paragraphs),
                "tableCount": len(document.tables),
            },
        }

    @classmethod
    def _extract_xlsx(cls, path: Path) -> dict[str, Any]:
        try:
            import openpyxl
        except ImportError as exc:
            return cls._unsupported_optional(".xlsx", exc)

        sheet_limit = ChatAttachmentContentService.file_extraction_xlsx_max_sheets()
        row_limit = ChatAttachmentContentService.file_extraction_xlsx_max_rows_per_sheet()

        try:
            workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
            lines: list[str] = []

            for sheet in workbook.worksheets[:sheet_limit]:
                lines.append(f"# Planilha: {sheet.title}")

                for row_index, row in enumerate(sheet.iter_rows(values_only=True)):
                    if row_index >= row_limit:
                        break

                    values = ["" if value is None else str(value) for value in row]

                    if any(values):
                        lines.append(" | ".join(values))

            content = "\n".join(lines).strip()

            if not content:
                return cls._unsupported_optional(
                    ".xlsx",
                    ValueError("empty_xlsx_content"),
                )

            return {
                "supported": True,
                "content": content,
                "metadata": {
                    "extractor": "openpyxl",
                    "extension": ".xlsx",
                    "sheetLimit": sheet_limit,
                    "rowLimit": row_limit,
                },
            }
        except Exception as exc:
            return cls._unsupported_optional(".xlsx", exc)

    @classmethod
    def _extract_xls(cls, path: Path) -> dict[str, Any]:
        try:
            import xlrd
        except ImportError as exc:
            return cls._unsupported_optional(".xls", exc)

        sheet_limit = ChatAttachmentContentService.file_extraction_xlsx_max_sheets()
        row_limit = ChatAttachmentContentService.file_extraction_xlsx_max_rows_per_sheet()

        try:
            workbook = xlrd.open_workbook(str(path))
            lines: list[str] = []

            for sheet in workbook.sheets()[:sheet_limit]:
                lines.append(f"# Planilha: {sheet.name}")

                for row_index in range(min(sheet.nrows, row_limit)):
                    values = [
                        "" if sheet.cell_value(row_index, col_index) in ("", None) else str(
                            sheet.cell_value(row_index, col_index)
                        )
                        for col_index in range(sheet.ncols)
                    ]

                    if any(values):
                        lines.append(" | ".join(values))

            content = "\n".join(lines).strip()

            if not content:
                return cls._legacy_office_format(".xls")

            return {
                "supported": True,
                "content": content,
                "metadata": {
                    "extractor": "xlrd",
                    "extension": ".xls",
                    "sheetLimit": sheet_limit,
                    "rowLimit": row_limit,
                },
            }
        except Exception:
            return cls._legacy_office_format(".xls")

    @classmethod
    def _extract_doc(cls, path: Path) -> dict[str, Any]:
        if not shutil.which("antiword"):
            return cls._legacy_office_format(".doc")

        timeout = ChatAttachmentContentService.file_extraction_subprocess_timeout_seconds()

        try:
            result = subprocess.run(
                ["antiword", "-m", "UTF-8.txt", str(path)],
                capture_output=True,
                text=True,
                timeout=timeout,
                check=False,
            )
        except (OSError, subprocess.TimeoutExpired):
            return cls._legacy_office_format(".doc")

        content = str(result.stdout or "").strip()

        if result.returncode != 0 or not content:
            return cls._legacy_office_format(".doc")

        return {
            "supported": True,
            "content": content,
            "metadata": {
                "extractor": "antiword",
                "extension": ".doc",
            },
        }

    @classmethod
    def _extract_pdf(cls, path: Path, *, page_limit: int | None) -> dict[str, Any]:
        from app.domain.services.chat_pdf_document_extraction_service import (
            ChatPdfDocumentExtractionService,
        )

        return ChatPdfDocumentExtractionService.extract_for_attachment_index(
            str(path),
            filename=path.name,
            page_limit=page_limit,
        )

    @classmethod
    def _extract_image(cls, path: Path, extension: str) -> dict[str, Any]:
        width: int | None = None
        height: int | None = None
        image_format: str | None = extension.lstrip(".").upper()

        try:
            from PIL import Image

            with Image.open(path) as image:
                width, height = image.size
                image_format = str(image.format or image_format or "").upper() or image_format
        except Exception:
            pass

        label = path.name
        descriptor = f"Imagem {label}"

        if width and height:
            descriptor = f"{descriptor} ({width}×{height}"

            if image_format:
                descriptor = f"{descriptor} {image_format}"

            descriptor = f"{descriptor})"

        ocr = ChatAttachmentImageOcrService.try_extract_text(path)
        ocr_text = str(ocr.get("text") or "").strip()

        metadata: dict[str, Any] = {
            "extension": extension,
            "width": width,
            "height": height,
            "format": image_format,
            "ocr": {
                "enabled": ChatAttachmentImageOcrService.is_enabled(),
                "used": bool(ocr.get("used")),
                "reason": ocr.get("reason"),
            },
        }

        if ocr_text:
            metadata["extractor"] = "image_ocr"
            metadata["ocr"]["charCount"] = ocr.get("charCount")
            content = (
                f"[{descriptor}]\n\n"
                "Texto extraído da imagem (OCR):\n"
                f"{ocr_text}"
            )
        else:
            metadata["extractor"] = "image_metadata"
            content = (
                f"[{descriptor}. "
                "Conteúdo visual indexado por metadados; descreva o que precisa "
                "(ex.: ler gráfico, extrair texto, gerar descrição ou texto alternativo).]"
            )

        return {
            "supported": True,
            "content": content,
            "metadata": metadata,
        }

    @classmethod
    def _legacy_office_format(cls, extension: str) -> dict[str, Any]:
        legacy = ChatAttachmentContentService.file_extraction_legacy_format(extension)
        reason = legacy.get("reason") or f"legacy_{extension.lstrip('.')}_format"
        user_hint = legacy.get("userHint") or ""

        return {
            "supported": False,
            "content": "",
            "metadata": {
                "reason": reason,
                "extension": extension,
                "userHint": user_hint,
            },
        }

    @classmethod
    def _unsupported_optional(cls, extension: str, exc: Exception) -> dict[str, Any]:
        return {
            "supported": False,
            "content": "",
            "metadata": {
                "reason": "optional_extractor_unavailable_or_failed",
                "extension": extension,
                "error": exc.__class__.__name__,
            },
        }
