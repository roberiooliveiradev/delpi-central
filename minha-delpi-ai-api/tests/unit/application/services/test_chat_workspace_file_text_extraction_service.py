from pathlib import Path

from app.application.services.chat_workspace_file_text_extraction_service import (
    ChatWorkspaceFileTextExtractionService,
)


def test_extract_docx_paragraphs_and_tables(tmp_path):
    from docx import Document

    path = Path(tmp_path) / "relatorio.docx"
    document = Document()
    document.add_paragraph("Título do relatório")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Produto"
    table.cell(0, 1).text = "Qtd"
    table.cell(1, 0).text = "10080001"
    table.cell(1, 1).text = "10"
    document.save(path)

    result = ChatWorkspaceFileTextExtractionService.extract(
        storage_path=str(path),
        filename="relatorio.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert result["supported"] is True
    assert result["metadata"]["extractor"] == "python_docx"
    assert "Título do relatório" in result["content"]
    assert "10080001" in result["content"]
    assert "Qtd" in result["content"]


def test_extract_xlsx_multiple_sheets(tmp_path):
    import openpyxl

    path = Path(tmp_path) / "dados.xlsx"
    workbook = openpyxl.Workbook()
    sheet_a = workbook.active
    sheet_a.title = "Resumo"
    sheet_a.append(["Filial", "Valor"])
    sheet_a.append(["01", "120"])
    sheet_b = workbook.create_sheet("Detalhe")
    sheet_b.append(["Item", "Qtd"])
    sheet_b.append(["Parafuso", "5"])
    workbook.save(path)

    result = ChatWorkspaceFileTextExtractionService.extract(
        storage_path=str(path),
        filename="dados.xlsx",
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    assert result["supported"] is True
    assert result["metadata"]["extractor"] == "openpyxl"
    assert "# Planilha: Resumo" in result["content"]
    assert "# Planilha: Detalhe" in result["content"]
    assert "Parafuso" in result["content"]


def test_extract_legacy_doc_without_antiword_returns_hint(tmp_path, monkeypatch):
    monkeypatch.setattr(
        "app.application.services.chat_workspace_file_text_extraction_service.shutil.which",
        lambda _cmd: None,
    )
    path = Path(tmp_path) / "ata.doc"
    path.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1")

    result = ChatWorkspaceFileTextExtractionService.extract(
        storage_path=str(path),
        filename="ata.doc",
        content_type="application/msword",
    )

    assert result["supported"] is False
    assert result["metadata"]["reason"] == "legacy_doc_format"
    assert "docx" in result["metadata"]["userHint"].lower()


def test_extract_doc_with_antiword(tmp_path, monkeypatch):
    path = Path(tmp_path) / "ata.doc"
    path.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1")

    monkeypatch.setattr(
        "app.application.services.chat_workspace_file_text_extraction_service.shutil.which",
        lambda cmd: "/usr/bin/antiword" if cmd == "antiword" else None,
    )

    class _Result:
        returncode = 0
        stdout = "Ata de reunião\nItem 1"

    monkeypatch.setattr(
        "app.application.services.chat_workspace_file_text_extraction_service.subprocess.run",
        lambda *args, **kwargs: _Result(),
    )

    result = ChatWorkspaceFileTextExtractionService.extract(
        storage_path=str(path),
        filename="ata.doc",
        content_type="application/msword",
    )

    assert result["supported"] is True
    assert result["metadata"]["extractor"] == "antiword"
    assert "Ata de reunião" in result["content"]


def test_extract_xls_reads_sheets(tmp_path, monkeypatch):
    path = Path(tmp_path) / "legacy.xls"
    path.write_bytes(b"fake-xls")

    class _Sheet:
        name = "Resumo"
        nrows = 2
        ncols = 2

        def cell_value(self, row: int, col: int) -> str:
            return [["Filial", "Valor"], ["01", "120"]][row][col]

    class _Workbook:
        def sheets(self):
            return [_Sheet()]

    monkeypatch.setattr(
        "xlrd.open_workbook",
        lambda _path: _Workbook(),
    )

    result = ChatWorkspaceFileTextExtractionService.extract(
        storage_path=str(path),
        filename="legacy.xls",
        content_type="application/vnd.ms-excel",
    )

    assert result["supported"] is True
    assert result["metadata"]["extractor"] == "xlrd"
    assert "# Planilha: Resumo" in result["content"]
    assert "120" in result["content"]
